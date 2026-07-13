from __future__ import annotations

import re
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree

import fitz
import pandas as pd
from docx import Document


SUPPORTED_SUFFIXES = {".docx", ".doc", ".pdf", ".txt", ".md", ".xlsx", ".xls", ".csv"}


def save_uploaded_file(uploaded_file, upload_dir: Path) -> Path:
    upload_dir.mkdir(parents=True, exist_ok=True)
    path = upload_dir / uploaded_file.name
    path.write_bytes(uploaded_file.getbuffer())
    return path


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".doc":
        return extract_legacy_doc_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix in {".xlsx", ".xls"}:
        return extract_excel_text(path)
    if suffix == ".csv":
        return pd.read_csv(path).to_csv(index=False)
    return ""


def extract_many(paths: Iterable[Path]) -> str:
    chunks = []
    for path in paths:
        text = extract_text(path)
        if text.strip():
            chunks.append(f"【文件：{path.name}】\n{text.strip()}")
    return "\n\n".join(chunks)


def extract_docx_text(path: Path) -> str:
    parts: list[str] = []
    try:
        doc = Document(path)
        for paragraph in doc.paragraphs:
            value = paragraph.text.strip()
            if value:
                parts.append(value)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
    except Exception:
        parts.append(_extract_docx_xml_text(path))
    return normalize_text("\n".join(parts))


def _extract_docx_xml_text(path: Path) -> str:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    text_nodes = [node.text or "" for node in root.findall(".//w:t", ns)]
    return "\n".join(text_nodes)


def extract_legacy_doc_text(path: Path) -> str:
    # .doc is a binary format. This best-effort extractor is enough for preview
    # and rule matching; users can still open/convert templates with WPS.
    data = path.read_bytes()
    candidates = []
    for encoding in ("utf-16le", "gb18030", "utf-8", "latin1"):
        try:
            candidates.append(data.decode(encoding, errors="ignore"))
        except Exception:
            continue
    text = "\n".join(candidates)
    pattern = r"[\u4e00-\u9fffA-Za-z0-9（）()《》、，。：；:;\-_/.\s]{6,}"
    matches = re.findall(pattern, text)
    return normalize_text("\n".join(m.strip() for m in matches if m.strip()))


def extract_pdf_text(path: Path) -> str:
    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text())
    return normalize_text("\n".join(parts))


def extract_excel_text(path: Path) -> str:
    sheets = pd.read_excel(path, sheet_name=None)
    chunks = []
    for name, df in sheets.items():
        chunks.append(f"Sheet: {name}")
        chunks.append(df.fillna("").to_csv(index=False))
    return normalize_text("\n".join(chunks))


def normalize_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def preview_text(text: str, limit: int = 4000) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n……（内容较长，已截断预览）"
