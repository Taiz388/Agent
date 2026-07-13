from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .utils import OUTPUTS_DIR, safe_filename


SELLER_INFO = {
    "seller_name": "福建省南铝板带加工有限公司",
    "seller_address": "南平市延平区水东街道工业路487号",
    "seller_bank": "中行南平分行",
    "seller_account": "426058390437",
    "seller_tax_no": "913507007821750903",
}

FONT_NAME = "宋体"
BLACK = RGBColor(0, 0, 0)
FALLBACK = "以双方确认资料为准"


def generate_contract_docx(
    contract_type: dict[str, Any],
    fields: dict[str, Any],
    products: list[dict[str, Any]],
    ai_clause: str = "",
) -> Path:
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)

    title = contract_type.get("title") or contract_type.get("name")
    add_title(doc, title)
    add_header_info(doc, contract_type, fields)
    add_parties(doc, contract_type, fields)
    add_intro(doc, contract_type, fields)
    add_product_table(doc, contract_type, products)
    add_contract_clauses(doc, contract_type, fields)
    add_standard_legal_safeguards(doc, contract_type)
    if ai_clause.strip():
        add_ai_clause(doc, ai_clause)
    add_contact_and_bank(doc, contract_type, fields)
    add_signature(doc, contract_type, fields)
    enforce_black_songti(doc)

    contract_no = fields.get("contract_no") or datetime.now().strftime("%Y%m%d%H%M")
    buyer = fields.get("buyer_name") or "客户"
    filename = safe_filename(f"{contract_no}_{buyer}_{contract_type.get('name')}.docx")
    output_path = OUTPUTS_DIR / filename
    doc.save(output_path)
    return output_path


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    for style_name, size, bold in [("Normal", 10.5, False), ("Heading 1", 12, True)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = BLACK


def add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(title or "合同"))
    run.bold = True
    run.font.size = Pt(18)


def add_header_info(doc: Document, contract_type: dict[str, Any], fields: dict[str, Any]) -> None:
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    values = [
        ("合同编号", field_value(fields.get("contract_no"))),
        ("签订日期", field_value(fields.get("sign_date"))),
        ("签订地点", field_value(fields.get("sign_place"), "南平市延平区")),
        ("合同类型", contract_type.get("name", "")),
    ]
    for idx, (label, value) in enumerate(values):
        row = idx // 2
        col = (idx % 2) * 2
        table.cell(row, col).text = label
        table.cell(row, col + 1).text = str(value)


def add_parties(doc: Document, contract_type: dict[str, Any], fields: dict[str, Any]) -> None:
    parties = contract_type.get("parties", {})
    buyer_label = parties.get("buyer_label", "需方")
    seller_label = parties.get("seller_label", "供方")
    seller = fields.get("seller_name") or parties.get("seller_default") or SELLER_INFO["seller_name"]
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = buyer_label
    table.cell(0, 1).text = field_value(fields.get("buyer_name"))
    table.cell(1, 0).text = seller_label
    table.cell(1, 1).text = field_value(seller)


def add_intro(doc: Document, contract_type: dict[str, Any], fields: dict[str, Any]) -> None:
    buyer = field_value(fields.get("buyer_name"), "需方")
    seller = field_value(fields.get("seller_name") or contract_type.get("parties", {}).get("seller_default") or SELLER_INFO["seller_name"])
    product = field_value(fields.get("product_name"), "合同标的物")
    project = fields.get("project_name")
    name = contract_type.get("name", "")
    group = contract_type.get("group", "")
    legal_basis = "依据《中华人民共和国民法典》《中华人民共和国产品质量法》及相关法律法规"
    if "出口" in name or "出口" in group:
        text = f"买卖双方经友好协商，就{buyer}向{seller}购买{product}事宜达成本合同；双方同意按本合同及双方确认文件履行。"
    elif contract_type.get("name") == "采购买卖合同":
        if project:
            text = f"经双方友好协商，{buyer}因{project}需要，向{seller}采购{product}，双方{legal_basis}订立本合同，共同遵守。"
        else:
            text = f"经双方友好协商，{buyer}向{seller}采购{product}，双方{legal_basis}订立本合同，共同遵守。"
    else:
        text = f"供需双方本着平等自愿、诚实信用、互惠互利的原则，就{buyer}向{seller}订购{product}事宜，{legal_basis}订立本合同，共同遵守。"
    doc.add_paragraph(text)


def add_product_table(doc: Document, contract_type: dict[str, Any], products: list[dict[str, Any]]) -> None:
    doc.add_heading("一、产品规格及价格", level=1)
    columns = contract_type.get("product_columns") or ["product_name", "specification", "quantity", "unit_price", "amount", "remark"]
    labels = {
        "item_no": "序号",
        "product_name": "产品",
        "specification": "规格",
        "alloy_state": "合金状态",
        "brand": "品牌",
        "unit": "单位",
        "quantity": "数量",
        "weight": "重量",
        "unit_price": "单价",
        "amount": "金额",
        "remark": "备注",
        "processing_fee": "加工费",
        "aluminum_price": "铝锭价",
        "po_no": "采购单号",
        "width_range": "宽度范围",
        "thickness_range": "厚度范围",
        "length_range": "长度范围",
        "color": "颜色",
        "front_coating": "正面涂层",
        "back_coating": "背面涂层",
        "coating": "涂层",
        "panel_type": "板型",
    }
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, col in enumerate(columns):
        table.rows[0].cells[idx].text = labels.get(col, col)
    for row_index, product in enumerate(products or [{}], start=1):
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = str(row_index) if col == "item_no" else field_value(product.get(col))
            cells[idx].text = value

    doc.add_paragraph("备注：最终货款金额以实际供货数量、双方确认单价及结算规则为准；未列明事项以双方书面确认的订单、技术协议或补充协议为准。")


def add_contract_clauses(doc: Document, contract_type: dict[str, Any], fields: dict[str, Any]) -> None:
    for index, clause in enumerate(contract_type.get("clauses", []), start=2):
        doc.add_heading(f"{to_chinese_index(index)}、{clause.get('heading', '')}", level=1)
        body = clause.get("body", "")
        body = fill_clause_body(body, fields)
        doc.add_paragraph(body)

    extra_fields = [
        ("交货地点", fields.get("delivery_place")),
        ("交货时间", fields.get("delivery_time")),
        ("付款方式", fields.get("payment_terms")),
        ("年提货量", fields.get("annual_quantity")),
        ("销售区域", fields.get("sales_region")),
        ("合同金额", fields.get("total_amount")),
    ]
    doc.add_heading("补充商务信息", level=1)
    for label, value in extra_fields:
        doc.add_paragraph(f"{label}：{field_value(value)}")


def add_standard_legal_safeguards(doc: Document, contract_type: dict[str, Any]) -> None:
    existing = "\n".join(clause.get("heading", "") for clause in contract_type.get("clauses", []))
    next_index = len(contract_type.get("clauses", [])) + 2
    standard = []
    if "不可抗力" not in existing:
        standard.append(("不可抗力", "不可抗力是指不能预见、不能避免且不能克服的客观情况。受不可抗力影响的一方应及时通知相对方，并在合理期限内提供证明，双方应采取必要措施减少损失；不可抗力影响消除后，应继续履行可履行部分。"))
    if "变更" not in existing and "补充" not in existing:
        standard.append(("合同变更与补充", "本合同未尽事宜，由双方另行协商并签署书面补充协议。补充协议、订单、技术协议、对账单等经双方确认的文件，与本合同具有同等效力。"))
    if "适用法律" not in existing and "争议" not in existing:
        standard.append(("适用法律与争议解决", "本合同订立、履行、解释及争议解决适用中华人民共和国法律。发生争议时，双方应先友好协商；协商不成的，按合同约定的管辖法院或争议解决方式处理。"))
    for offset, (heading, body) in enumerate(standard):
        doc.add_heading(f"{to_chinese_index(next_index + offset)}、{heading}", level=1)
        doc.add_paragraph(body)


def fill_clause_body(body: str, fields: dict[str, Any]) -> str:
    for key, value in fields.items():
        body = body.replace("{{ " + key + " }}", str(value))
        body = body.replace("{{" + key + "}}", str(value))
    return body


def add_ai_clause(doc: Document, ai_clause: str) -> None:
    doc.add_heading("AI 辅助补充条款草稿", level=1)
    for line in ai_clause.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())


def add_contact_and_bank(doc: Document, contract_type: dict[str, Any], fields: dict[str, Any]) -> None:
    doc.add_heading("通知、送达及账户信息", level=1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    rows = [
        ("项目", "需方/甲方", "供方/乙方"),
        ("单位名称", field_value(fields.get("buyer_name")), field_value(fields.get("seller_name") or SELLER_INFO["seller_name"])),
        ("单位地址", field_value(fields.get("buyer_address")), SELLER_INFO["seller_address"]),
        ("联系人", field_value(fields.get("contact_person")), field_value(fields.get("seller_contact"))),
        ("联系电话", field_value(fields.get("contact_phone")), field_value(fields.get("seller_phone"))),
        ("开户行/账号/税号", field_value(fields.get("buyer_bank_info")), f"{SELLER_INFO['seller_bank']} / {SELLER_INFO['seller_account']} / {SELLER_INFO['seller_tax_no']}"),
    ]
    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            table.cell(row_idx, col_idx).text = str(value)


def add_signature(doc: Document, contract_type: dict[str, Any], fields: dict[str, Any]) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    left_label = contract_type.get("parties", {}).get("buyer_label", "需方")
    right_label = contract_type.get("parties", {}).get("seller_label", "供方")
    rows = [
        (f"{left_label}（盖章）：", f"{right_label}（盖章）："),
        ("法定代表人或授权代表：", "法定代表人或授权代表："),
        ("日期：", "日期："),
        ("", ""),
    ]
    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            table.cell(row_idx, col_idx).text = value


def field_value(value: Any, fallback: str = FALLBACK) -> str:
    text = str(value or "").strip()
    return text if text else fallback


def enforce_black_songti(doc: Document) -> None:
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
            if not run.font.name:
                run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK
                        if not run.font.name:
                            run.font.name = FONT_NAME
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def to_chinese_index(value: int) -> str:
    mapping = {
        1: "一",
        2: "二",
        3: "三",
        4: "四",
        5: "五",
        6: "六",
        7: "七",
        8: "八",
        9: "九",
        10: "十",
        11: "十一",
        12: "十二",
        13: "十三",
        14: "十四",
        15: "十五",
        16: "十六",
        17: "十七",
        18: "十八",
    }
    return mapping.get(value, str(value))
