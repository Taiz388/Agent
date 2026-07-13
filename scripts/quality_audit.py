from __future__ import annotations

import sys
from pathlib import Path
from typing import Iterable

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document

from src.auto_contract import build_fields_from_text, contract_type_for_key, detect_contract_key
from src.document_parser import extract_many
from src.generator import generate_contract_docx
from src.utils import load_contract_config

REFERENCE_DIR = ROOT / "参考资料"
BLACK = "000000"

CASES = [
    {
        "name": "标书资料-16#挤压筒换内衬",
        "mode": "标书资料生成合同",
        "path": REFERENCE_DIR / "交易生成合同" / "MLBB（20）24-29616#挤压筒换内衬",
        "expected": ["MLBB（20）24-296", "福建省南平铝业股份有限公司", "江苏xxx集团股份有限公司", "16#挤压筒换内衬", "285000", "违约责任", "争议解决", "不可抗力"],
        "min_files": 3,
    },
    {
        "name": "标书资料-电解槽用纳米隔热板采购",
        "mode": "标书资料生成合同",
        "path": REFERENCE_DIR / "交易生成合同" / "NG（32）22-227电解槽用纳米隔热板采购",
        "expected": ["NG（32）22-227", "福建省南平铝业股份有限公司", "天津xxxx有限公司", "电解槽用纳米隔热板", "597000", "增值税专用发票", "违约责任", "争议解决"],
        "min_files": 3,
    },
    {
        "name": "模板套用-板带零售合同",
        "mode": "套用现有模板",
        "path": REFERENCE_DIR / "20260707南铝板带销售合同及模板" / "20260707南铝板带销售合同及模板" / "已签合同" / "带材零售合同  BX（LS）26-040【浙江银环】.docx",
        "expected": ["BX（LS）26-040", "铝板", "浙", "福建省南铝", "民法典", "产品质量法", "违约责任", "争议"],
        "min_files": 1,
    },
    {
        "name": "模板套用-出口销售合同",
        "mode": "套用现有模板",
        "path": REFERENCE_DIR / "20260707南铝板带销售合同及模板" / "20260707南铝板带销售合同及模板" / "已签合同" / "BX(CK)26-011南铝板带出口合同.docx",
        "expected": ["SALES CONTRACT", "BX(CK)26-011", "KP", "FUJIAN NANPING", "Commodity", "Payment", "Dispute"],
        "min_files": 1,
    },
]


def case_files(path: Path) -> list[Path]:
    if path.is_file():
        return [path]
    return [p for p in sorted(path.iterdir()) if p.is_file() and p.suffix.lower() in {".doc", ".docx", ".pdf", ".xlsx", ".xls", ".txt", ".md"}]


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def iter_runs(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            yield run
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        yield run


def assert_black_font(path: Path) -> list[str]:
    doc = Document(path)
    bad = []
    for run in iter_runs(doc):
        if not run.text.strip():
            continue
        rgb = run.font.color.rgb
        if rgb is None or str(rgb).upper() != BLACK:
            bad.append(run.text[:20])
            if len(bad) >= 5:
                break
    return bad


def run_case(case: dict) -> tuple[bool, str]:
    config = load_contract_config()
    paths = case_files(case["path"])
    if len(paths) < case.get("min_files", 1):
        return False, f"上传文件数量不足：{len(paths)}"
    text = extract_many(paths)
    key = detect_contract_key(case["mode"], text, config)
    contract_type = contract_type_for_key(key, config)
    fields = build_fields_from_text(text, contract_type, case["mode"])
    output_path = generate_contract_docx(contract_type, fields, fields.get("products", []), ai_clause="")
    generated = docx_text(output_path)

    missing = [item for item in case["expected"] if item not in generated]
    garbled = [item for item in ["鍚", "閾", "绂", "�", "????"] if item in generated]
    color_bad = assert_black_font(output_path)
    if missing or garbled or color_bad:
        details = []
        if missing:
            details.append(f"缺少：{missing}")
        if garbled:
            details.append(f"疑似乱码：{garbled}")
        if color_bad:
            details.append(f"非显式黑色字体片段：{color_bad}")
        return False, "；".join(details) + f"；输出：{output_path.name}"
    return True, f"通过；输出：{output_path.name}"


def main() -> None:
    failures = []
    for case in CASES:
        ok, message = run_case(case)
        print(f"[{'PASS' if ok else 'FAIL'}] {case['name']} - {message}")
        if not ok:
            failures.append(case["name"])
    if failures:
        raise SystemExit("质量验收失败：" + "、".join(failures))


if __name__ == "__main__":
    main()
