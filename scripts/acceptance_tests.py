from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document

from src.auto_contract import build_fields_from_text, contract_type_for_key, detect_contract_key
from src.document_parser import extract_many
from src.generator import generate_contract_docx
from src.utils import load_contract_config

REFERENCE_DIR = ROOT / "参考资料"

CASES = [
    {
        "name": "交易案例-16#挤压筒换内衬",
        "mode": "标书资料生成合同",
        "path": REFERENCE_DIR / "交易生成合同" / "MLBB（20）24-29616#挤压筒换内衬",
        "contains": ["16#挤压筒换内衬", "江苏xxx集团股份有限公司", "285000"],
    },
    {
        "name": "交易案例-电解槽用纳米隔热板采购",
        "mode": "标书资料生成合同",
        "path": REFERENCE_DIR / "交易生成合同" / "NG（32）22-227电解槽用纳米隔热板采购",
        "contains": ["电解槽用纳米隔热板", "天津xxxx有限公司", "597000"],
    },
    {
        "name": "模板案例-板带零售合同",
        "mode": "套用现有模板",
        "path": REFERENCE_DIR / "20260707南铝板带销售合同及模板" / "20260707南铝板带销售合同及模板" / "已签合同" / "带材零售合同  BX（LS）26-040【浙江银环】.docx",
        "contains": ["铝板", "BX（LS）26-040", "浙江"],
    },
    {
        "name": "模板案例-出口销售合同",
        "mode": "套用现有模板",
        "path": REFERENCE_DIR / "20260707南铝板带销售合同及模板" / "20260707南铝板带销售合同及模板" / "已签合同" / "BX(CK)26-011南铝板带出口合同.docx",
        "contains": ["SALES CONTRACT", "BX(CK)26-011", "KP"],
    },
]


def case_files(path: Path) -> list[Path]:
    if path.is_file():
        return [path]
    return [p for p in sorted(path.iterdir()) if p.is_file() and p.suffix.lower() in {".doc", ".docx", ".pdf", ".xlsx"}]


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def run_case(case: dict) -> tuple[bool, str]:
    config = load_contract_config()
    paths = case_files(case["path"])
    if not paths:
        return False, "未找到测试文件"
    text = extract_many(paths)
    key = detect_contract_key(case["mode"], text, config)
    contract_type = contract_type_for_key(key, config)
    fields = build_fields_from_text(text, contract_type, case["mode"])
    output_path = generate_contract_docx(contract_type, fields, fields.get("products", []), ai_clause="")
    generated = docx_text(output_path)
    missing = [item for item in case["contains"] if item not in generated]
    if missing:
        return False, f"缺少预期内容：{missing}；输出：{output_path.name}"
    return True, f"通过；输出：{output_path.name}"


def main() -> None:
    failures = []
    for case in CASES:
        ok, message = run_case(case)
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {case['name']} - {message}")
        if not ok:
            failures.append(case["name"])
    if failures:
        raise SystemExit("验收失败：" + "、".join(failures))


if __name__ == "__main__":
    main()
